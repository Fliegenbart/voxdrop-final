"""
DOCX Parser - Extract text, structure, and metadata from Word documents.
Uses python-docx for parsing.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from typing import Dict, Any, List


async def parse_docx(file_path: str) -> Dict[str, Any]:
    """
    Parse a DOCX file and extract relevant content.

    Returns:
        dict with:
        - text: Full text content
        - paragraphs: List of paragraph contents
        - headings: Detected headings
        - images: Image information
        - colors: Detected colors
        - fonts: Font information
        - metadata: Document metadata
    """
    doc = Document(file_path)

    full_text = ""
    paragraphs = []
    headings = []
    images = []
    colors = set()
    fonts = set()
    lists = []

    # Track heading levels
    heading_styles = {
        "Heading 1": 1,
        "Heading 2": 2,
        "Heading 3": 3,
        "Heading 4": 4,
        "Title": 1,
    }

    for para in doc.paragraphs:
        para_text = para.text.strip()

        # Check if it's a heading
        if para.style and para.style.name in heading_styles:
            if para_text:
                headings.append({
                    "text": para_text,
                    "level": heading_styles[para.style.name],
                    "style": para.style.name,
                })

        # Check if it's a list item
        is_list_item = False
        if para.style and ("List" in para.style.name or "Bullet" in para.style.name):
            is_list_item = True
            lists.append(para_text)

        # Extract font and color information from runs
        for run in para.runs:
            if run.font:
                font_name = run.font.name or "Default"
                font_size = run.font.size.pt if run.font.size else 12
                fonts.add((font_name, font_size))

                # Get text color
                if run.font.color and run.font.color.rgb:
                    rgb = run.font.color.rgb
                    # RGBColor stores as hex string
                    r = int(str(rgb)[0:2], 16)
                    g = int(str(rgb)[2:4], 16)
                    b = int(str(rgb)[4:6], 16)
                    colors.add((r, g, b))

        para_info = {
            "text": para_text,
            "style": para.style.name if para.style else "Normal",
            "is_heading": para.style.name in heading_styles if para.style else False,
            "is_list_item": is_list_item,
            "alignment": str(para.alignment) if para.alignment else "left",
        }
        paragraphs.append(para_info)
        full_text += para_text + "\n"

    # Count images (inline shapes)
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_count += 1
            images.append({
                "type": "inline",
                "has_alt_text": False,  # Would need more complex parsing
            })

    # Extract tables
    tables = []
    for table in doc.tables:
        table_data = {
            "rows": len(table.rows),
            "cols": len(table.columns),
            "has_header": True,  # Assume first row is header
            "content": [],
        }
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            table_data["content"].append(row_text)
            full_text += " | ".join(row_text) + "\n"
        tables.append(table_data)

    # Get core properties (metadata)
    metadata = {}
    try:
        core_props = doc.core_properties
        metadata = {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
        }
    except Exception:
        metadata = {"title": "", "author": "", "subject": ""}

    return {
        "text": full_text.strip(),
        "paragraphs": paragraphs,
        "paragraph_count": len(paragraphs),
        "headings": headings,
        "images": images,
        "image_count": len(images),
        "tables": tables,
        "table_count": len(tables),
        "lists": lists,
        "list_count": len(lists),
        "colors": list(colors),
        "fonts": list(fonts),
        "metadata": metadata,
        "format": "docx",
    }
